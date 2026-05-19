import streamlit as st
import docx
import io

# 1. HIGH-CONTRAST DYNAMIC COLOR STYLING
st.set_page_config(page_title="AI Student Report Hub", page_icon="🎓", layout="centered")

st.markdown("""
    <style>
    /* Gradient Pastel Background */
    .stApp {
        background: linear-gradient(135deg, #e0f2fe 0%, #f3e8ff 50%, #fce7f3 100%);
    }
    
    /* Vibrant Header Title styling */
    h1 { 
        color: #701a75 !important; 
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-weight: 800;
        text-align: center;
        padding-bottom: 10px;
    }
    
    /* Neon Highlighted Inputs */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea {
        border-radius: 12px !important;
        border: 2px solid #a855f7 !important;
        background-color: #ffffff !important;
        font-size: 16px !important;
    }
    
    /* Electric Orange Action Button */
    .stButton>button {
        background: linear-gradient(90deg, #ea580c 0%, #db2777 100%) !important;
        color: white !important;
        border-radius: 15px !important;
        height: 3.5em !important;
        width: 100% !important;
        border: none !important;
        font-size: 18px !important;
        font-weight: bold !important;
        box-shadow: 0 4px 12px rgba(219, 39, 119, 0.2);
    }
    
    .stButton>button:hover {
        background: linear-gradient(90deg, #db2777 0%, #ea580c 100%) !important;
    }
    
    /* File Uploader Container */
    section[data-testid="stFileUploadDropzone"] {
        border-radius: 15px !important;
        border: 2px dashed #db2777 !important;
        background-color: rgba(255, 255, 255, 0.6) !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("🎓 AI Student Report Hub")
st.write("<p style='text-align: center; color: #6b21a8; font-weight: bold;'>Format, structure, and customize your academic documentation seamlessly.</p>", unsafe_allow_html=True)

# 2. STUDENT INFORMATION AREA
col1, col2 = st.columns(2)
with col1:
    student_name = st.text_input("👤 Student Name", placeholder="e.g., Bolla Anusha")
with col2:
    roll_number = st.text_input("🆔 Roll Number / USN", placeholder="e.g., 1VT23CS020")

report_topic = st.text_input("📚 Enter Your Report Topic", placeholder="e.g., Blockchain Waste Management")

st.markdown("<hr style='border: 1px solid #c084fc;'>", unsafe_allow_html=True)
st.markdown("<h3 style='color: #701a75;'>🛠️ Structure & Content Adaption</h3>", unsafe_allow_html=True)

# THE TYPE BAR
user_description = st.text_area(
    "💬 Type Bar: Describe exactly how the report should be generated:", 
    placeholder="e.g., Keep headings strictly aligned with standard VTU criteria, emphasize system architecture blocks..."
)

# THE UPLOAD OPTION
uploaded_sample = st.file_uploader("📂 Upload Your Sample Report Document (.docx)", type=["docx"])

if uploaded_sample:
    st.success(f"✅ Layout Reference Sample Loaded: {uploaded_sample.name}")

# 3. ROBUST FILE PROCESSING LOGIC
if st.button("🚀 Generate Academic Report"):
    if student_name and report_topic:
        with st.spinner("⚡ Formatting your customized layout structure..."):
            
            output_doc = docx.Document()
            
            # Formulating the custom header layout
            output_doc.add_heading(f"Report: {report_topic.upper()}", level=0)
            author_p = output_doc.add_paragraph()
            author_p.add_run(f"Prepared by: {student_name}\n").bold = True
            author_p.add_run(f"Seat Number/USN: {roll_number}\n")
            if user_description:
                author_p.add_run(f"Design Directives: {user_description}").italic = True
                
            output_doc.add_page_break()
            
            if uploaded_sample is not None:
                try:
                    # Read structural lines safely out of the sample doc stream
                    sample_stream = docx.Document(uploaded_sample)
                    
                    for paragraph in sample_stream.paragraphs:
                        text_content = paragraph.text.strip()
                        if len(text_content) > 0:
                            # Replicate structural headings automatically
                            if paragraph.style.name.startswith('Heading'):
                                try:
                                    level_idx = int(paragraph.style.name[-1])
                                    output_doc.add_heading(text_content, level=level_idx)
                                except ValueError:
                                    output_doc.add_heading(text_content, level=1)
                            else:
                                # Apply internal contextual rewriting mechanics to clear out old subject text
                                if "introduction" in text_content.lower():
                                    output_doc.add_paragraph(f"This introductory analysis evaluates the engineering deployment of {report_topic}. This technical architecture resolves core data optimization bottlenecks, following the programmatic layout specified in structural student reference records.")
                                elif "conclusion" in text_content.lower() or "summary" in text_content.lower():
                                    output_doc.add_paragraph(f"In summary, implementing the components of {report_topic} ensures systemic performance control. Future scope milestones will continue optimizing operational throughput and structural security metrics.")
                                else:
                                    output_doc.add_paragraph(f"Regarding the technical execution metrics of {report_topic}, development cycles are organized to adapt seamlessly. Emphasizing the instructions processed via the user parameters ('{user_description}'), system constraints are structured to fulfill academic project parameters.")
                except Exception as parse_error:
                    st.error(f"Note on structure assembly: {parse_error}")
            else:
                # Clean fallback format block
                output_doc.add_heading("1. Introduction", level=1)
                output_doc.add_paragraph(f"This document presents a dedicated study of {report_topic} configured under the guidance criteria: {user_description}.")
                output_doc.add_heading("2. Core Implementation Strategy", level=1)
                output_doc.add_paragraph("Structural layout generated to comply with standard academic formatting specifications.")

            # File compilation down to byte buffer
            byte_stream = io.BytesIO()
            output_doc.save(byte_stream)
            byte_stream.seek(0)
            
            st.success("🎉 Success! Your document structure has been cleanly generated.")
            st.download_button(
                label="📥 Click Here to Download Your Custom Word Document",
                data=byte_stream.getvalue(),
                file_name=f"{report_topic.replace(' ', '_')}_Custom_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Please ensure Student Name and Report Topic fields are completely filled out first!")
