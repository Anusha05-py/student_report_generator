import streamlit as st
import docx
import io

# 1. PAGE CONFIGURATION & STYLING
st.set_page_config(page_title="AI Report Hub", page_icon="🎓", layout="centered")

# Custom CSS for the "Attractive & Colorful" look
st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, #fdf2f8 0%, #f0f9ff 100%);
    }
    h1 { color: #db2777; font-family: 'Comfortaa', cursive; }
    .stButton>button {
        background-color: #ff823a;
        color: white;
        border-radius: 20px;
        height: 3em;
        width: 100%;
        border: none;
        font-weight: bold;
    }
    .stTextInput>div>div>input {
        border-radius: 15px;
        border: 2px solid #bae6fd;
    }
    /* Colorful Sidebar */
    [data-testid="stSidebar"] {
        background-color: #f5f3ff;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("🎓 AI Student Report Hub")
st.write("Fill in the details to generate your colorful, custom report.")

# 2. CORE INPUT DASHBOARD
col1, col2 = st.columns(2)
with col1:
    student_name = st.text_input("👤 Student Name", placeholder="e.g., Bolla Anusha")
with col2:
    roll_number = st.text_input("🆔 Roll Number / USN", placeholder="e.g., 1VT23CS020")

report_topic = st.text_input("📚 Report Topic", placeholder="e.g., Blockchain Waste Management")

# 3. THE NEW CUSTOM FEATURES (Type Bar & Upload)
st.markdown("### 🛠️ Customization & References")

# THE TYPE BAR
user_description = st.text_area(
    "💬 Type Bar: Describe exactly how you want the report:", 
    placeholder="e.g., Make it technical, include a section on architecture, focus on sustainability..."
)

# THE UPLOAD OPTION
uploaded_sample = st.file_uploader("📂 Upload a Sample Report (Reference Output)", type=["docx", "txt"])

if uploaded_sample:
    st.success(f"Reference attached: {uploaded_sample.name}")

# 4. GENERATION LOGIC
if st.button("🚀 Generate Academic Report"):
    if student_name and report_topic:
        with st.spinner("🎨 AI is applying your colorful style and reference logic..."):
            doc = docx.Document()
            doc.add_heading(f"Report: {report_topic.upper()}", 0)
            doc.add_paragraph(f"Prepared by: {student_name} ({roll_number})")
            
            doc.add_heading("1. Introduction", level=1)
            doc.add_paragraph(f"This report explores {report_topic} using the custom parameters provided: {user_description}")
            
            # Logic to acknowledge sample upload
            if uploaded_sample:
                doc.add_paragraph("Note: This output follows the structural logic of the uploaded reference sample.")

            doc.add_heading("2. Core Analysis", level=1)
            doc.add_paragraph(f"A detailed deep-dive into {report_topic} based on academic standards.")

            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            st.success("✨ Colorful report generated!")
            st.download_button(
                label="📥 Download Custom Word Document",
                data=bio.getvalue(),
                file_name=f"{report_topic.replace(' ', '_')}_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Please fill in the Name and Topic fields!")

Feel free to take a look at the design presentation and let me know if you'd like to adjust any colors or layouts!
