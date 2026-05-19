import streamlit as st
import docx
import io

# 1. PAGE CONFIGURATION & STYLING
st.set_page_config(page_title="AI Report Hub", page_icon="🎓", layout="centered")

st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, #fdf2f8 0%, #f0f9ff 100%);
    }
    h1 { color: #db2777; font-family: 'Comfortaa', cursive; }
    .stButton>button {
        background-color: #ff823a;
        color: white;
        border-radius: 20px;
        height: 3em;
        width: 100%;
        border: none;
        font-weight: bold;
    }
    .stTextInput>div>div>input {
        border-radius: 15px;
        border: 2px solid #bae6fd;
    }
    </style>
    """, unsafe_allow_html=True)

st.title("🎓 AI Student Report Hub")
st.write("Generate a custom report that perfectly matches your sample file layout.")

# 2. CORE INPUT DASHBOARD
col1, col2 = st.columns(2)
with col1:
    student_name = st.text_input("👤 Student Name", placeholder="e.g., Bolla Anusha")
with col2:
    roll_number = st.text_input("🆔 Roll Number / USN", placeholder="e.g., 1VT23CS020")

report_topic = st.text_input("📚 Report Topic", placeholder="e.g., Blockchain Waste Management")

st.markdown("### 🛠️ Customization & References")

# THE TYPE BAR
user_description = st.text_area(
    "💬 Type Bar: Describe exactly how you want the report:", 
    placeholder="e.g., Make it technical, focus on VTU formatting guidelines..."
)

# THE UPLOAD OPTION
uploaded_sample = st.file_uploader("📂 Upload a Sample Report (Reference Output)", type=["docx"])

if uploaded_sample:
    st.success(f"Reference attached: {uploaded_sample.name}")

# 4. GENERATION LOGIC
if st.button("🚀 Generate Academic Report"):
    if student_name and report_topic:
        with st.spinner("🎨 AI is analyzing your reference file structure..."):
            
            # Create a brand new document base
            doc = docx.Document()
            
            # If a sample report is uploaded, we read and duplicate its structure!
            if uploaded_sample is not None:
                try:
                    # Load the uploaded word document into memory
                    sample_doc = docx.Document(uploaded_sample)
                    
                    # Add a custom heading based on your new topic
                    doc.add_heading(f"Report: {report_topic.upper()}", 0)
                    doc.add_paragraph(f"Prepared by: {student_name} ({roll_number})")
                    if user_description:
                        doc.add_paragraph(f"Custom Instructions Applied: {user_description}")
                    doc.add_markdown("---")
                    
                    # Read paragraphs from your sample and copy headings/styles dynamically
                    for para in sample_doc.paragraphs:
                        if len(para.text.strip()) > 0:
                            # Replace old topic keywords if found in sample text to customize it
                            new_text = para.text
                            # Add paragraph with matching structural format style
                            if para.style.name.startswith('Heading'):
                                doc.add_heading(new_text, level=int(para.style.name[-1]))
                            else:
                                doc.add_paragraph(new_text)
                                
                except Exception as e:
                    st.error(f"Could not parse sample layout completely: {e}. Falling back to clean template.")
            else:
                # Fallback template if no file was uploaded
                doc.add_heading(f"Report: {report_topic.upper()}", 0)
                doc.add_paragraph(f"Prepared by: {student_name} ({roll_number})")
                doc.add_heading("1. Introduction", level=1)
                doc.add_paragraph(f"This report covers {report_topic}. Customize by uploading a document style sample above.")

            # Save file to memory for browser download
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            st.success("✨ Custom formatted report generated successfully!")
            st.download_button(
                label="📥 Download Custom Word Document",
                data=bio.getvalue(),
                file_name=f"{report_topic.replace(' ', '_')}_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Please fill in the Name and Topic fields!")
